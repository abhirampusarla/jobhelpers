"""
Resume Maker
Tailors a base resume to a job URL and writes a DOCX resume.
"""

import argparse
import json
import logging
import re
import time
import zipfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree
from xml.sax.saxutils import escape

import requests
from bs4 import BeautifulSoup

import constants

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    Document = None
    Pt = None

logging.basicConfig(level=logging.INFO, format="%(levelname)s %(message)s")
log = logging.getLogger(__name__)


def load_prompts() -> dict[str, str]:
    prompts = {}
    current_key = None
    current_lines = []

    def flush_current():
        if current_key:
            prompts[current_key] = "\n".join(current_lines).strip()

    lines = constants.PROMPTS_FILE.read_text(encoding="utf-8").splitlines()
    for line in lines:
        if current_key and (line.startswith("  ") or not line.strip()):
            current_lines.append(line[2:] if line.startswith("  ") else "")
            continue

        flush_current()
        current_key = None
        current_lines = []

        stripped = line.strip()
        if not stripped or stripped.startswith("#"):
            continue
        if stripped.endswith(": |"):
            current_key = stripped[:-3].strip()
            continue
        if ":" in stripped:
            key, value = stripped.split(":", 1)
            prompts[key.strip()] = value.strip().strip('"')

    flush_current()
    required = {"tailor_resume_system", "tailor_resume_user"}
    missing = sorted(required - prompts.keys())
    if missing:
        raise SystemExit(f"prompts.yaml missing required prompt(s): {', '.join(missing)}")
    return prompts


def render_prompt(template: str, **values: object) -> str:
    rendered = template
    for key, value in values.items():
        rendered = rendered.replace(f"{{{key}}}", str(value))
    return rendered


def read_docx(path: Path) -> str:
    if Document is None:
        return read_docx_basic(path)

    document = Document(path)
    parts = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def read_docx_basic(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml")

    root = ElementTree.fromstring(xml)
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    paragraphs = []
    for paragraph in root.findall(".//w:p", namespace):
        texts = [node.text or "" for node in paragraph.findall(".//w:t", namespace)]
        text = "".join(texts).strip()
        if text:
            paragraphs.append(text)
    return "\n".join(paragraphs).strip()


def read_text_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8").strip()
    if suffix == ".docx":
        return read_docx(path)
    raise SystemExit(f"Unsupported file type: {path.suffix}. Use .txt, .md, or .docx")


def fetch_job_url(url: str) -> str:
    log.info(f"Fetching job URL: {url}")
    resp = requests.get(url, headers=constants.SESSION_HEADERS, timeout=constants.REQUEST_TIMEOUT)
    resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "noscript", "svg", "header", "footer", "nav"]):
        tag.decompose()

    text = soup.get_text(separator="\n", strip=True)
    lines = []
    seen = set()
    for line in text.splitlines():
        clean = " ".join(line.split())
        if len(clean) < 3 or clean in seen:
            continue
        seen.add(clean)
        lines.append(clean)

    return "\n".join(lines)[:constants.MAX_JOB_TEXT_CHARS].strip()


def load_job_description(job_url: str) -> str:
    return fetch_job_url(job_url)


def prompt_for_job_url() -> str:
    url = input("Paste job URL: ").strip()
    if not url:
        raise SystemExit("Job URL is required.")
    return url


def prompt_for_resume_path(default_path: Path) -> Path:
    print(f"Base resume not found at: {default_path}")
    raw_path = input("Paste base resume path (.docx, .txt, or .md): ").strip()
    if not raw_path:
        raise SystemExit("Base resume path is required.")
    return Path(raw_path).expanduser()


def resolve_resume_path(raw_path: str) -> Path:
    path = Path(raw_path).expanduser()
    if path.exists():
        return path
    prompted_path = prompt_for_resume_path(path)
    if not prompted_path.exists():
        raise SystemExit(f"Base resume not found: {prompted_path}")
    return prompted_path


def safe_filename_part(text: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_")
    return cleaned or "Tailored_Resume"


def infer_role_name(result: dict, job_description: str) -> str:
    resume = result.get("resume") if isinstance(result.get("resume"), dict) else {}
    for candidate in [
        result.get("target_role"),
        resume.get("headline"),
    ]:
        if candidate and str(candidate).strip():
            return safe_filename_part(str(candidate))[:80]

    role_keywords = (
        "analyst",
        "engineer",
        "scientist",
        "developer",
        "architect",
        "manager",
        "consultant",
        "specialist",
    )
    for line in job_description.splitlines()[:80]:
        clean = " ".join(line.split())
        if 4 <= len(clean) <= 100 and any(keyword in clean.lower() for keyword in role_keywords):
            return safe_filename_part(clean)[:80]
    return "Tailored_Resume"


def default_output_paths(result: dict, job_description: str) -> tuple[Path, Path]:
    role_name = infer_role_name(result, job_description)
    base_name = f"{constants.OUTPUT_NAME_PREFIX}_{role_name}"
    return (
        constants.OUTPUT_DIR / f"{base_name}.docx",
        constants.OUTPUT_DIR / f"{base_name}_analysis.json",
    )


def parse_json_object(text: str) -> dict:
    text = text.strip()
    if text.startswith("```"):
        text = text.strip("`")
        if text.lower().startswith("json"):
            text = text[4:].strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}")
        if start >= 0 and end > start:
            return json.loads(text[start:end + 1])
        raise


def resume_json_schema_text() -> str:
    return json.dumps({
        "target_role": "...",
        "resume": {
            "name": "...",
            "headline": "...",
            "contact": ["..."],
            "summary": ["..."],
            "skills": {"Category": ["skill"]},
            "experience": [
                {
                    "company": "...",
                    "title": "...",
                    "location": "...",
                    "dates": "...",
                    "bullets": ["..."],
                }
            ],
            "projects": [{"name": "...", "bullets": ["..."]}],
            "education": [{"school": "...", "degree": "...", "details": "..."}],
            "certifications": ["..."],
        },
        "fit_summary": "...",
        "keywords_used": ["..."],
        "gaps_or_risks": ["..."],
        "edit_notes": ["..."],
    }, indent=2)


def post_openai_with_retries(payload: dict) -> dict:
    last_error = None
    for attempt in range(1, constants.LLM_MAX_RETRIES + 1):
        try:
            resp = requests.post(
                constants.OPENAI_CHAT_COMPLETIONS_URL,
                headers={
                    "Authorization": f"Bearer {constants.OPENAI_API_KEY}",
                    "Content-Type": "application/json",
                },
                json=payload,
                timeout=constants.LLM_TIMEOUT,
            )
            if resp.status_code == 429:
                retry_after = resp.headers.get("retry-after")
                wait_seconds = int(retry_after) if retry_after and retry_after.isdigit() else constants.LLM_RETRY_SECONDS
                last_error = resp
                if attempt < constants.LLM_MAX_RETRIES:
                    log.warning(f"OpenAI rate limited the request. Retrying in {wait_seconds}s ({attempt}/{constants.LLM_MAX_RETRIES})...")
                    time.sleep(wait_seconds)
                    continue
            resp.raise_for_status()
            return resp.json()
        except requests.exceptions.RequestException as e:
            last_error = e.response if getattr(e, "response", None) is not None else e
            if attempt < constants.LLM_MAX_RETRIES:
                log.warning(f"OpenAI request failed. Retrying in {constants.LLM_RETRY_SECONDS}s ({attempt}/{constants.LLM_MAX_RETRIES}): {e}")
                time.sleep(constants.LLM_RETRY_SECONDS)
                continue
            break

    if getattr(last_error, "status_code", None) == 429:
        detail = ""
        try:
            detail = last_error.json().get("error", {}).get("message", "")
        except ValueError:
            detail = last_error.text[:300]
        raise SystemExit(
            "OpenAI returned 429 Too Many Requests after retries. "
            "This usually means rate limit or insufficient quota. "
            "Wait a minute and retry, or check billing/limits for the API key. "
            f"Details: {detail}"
        )

    raise SystemExit(f"OpenAI request failed after retries: {last_error}")


def post_claude_with_retries(payload: dict) -> dict:
    last_error = None
    for attempt in range(1, constants.LLM_MAX_RETRIES + 1):
        try:
            resp = requests.post(
                constants.ANTHROPIC_MESSAGES_URL,
                headers={
                    "x-api-key": constants.CLAUDE_API_KEY,
                    "anthropic-version": constants.ANTHROPIC_VERSION,
                    "content-type": "application/json",
                },
                json=payload,
                timeout=constants.LLM_TIMEOUT,
            )
            if resp.status_code == 429:
                retry_after = resp.headers.get("retry-after")
                wait_seconds = int(retry_after) if retry_after and retry_after.isdigit() else constants.LLM_RETRY_SECONDS
                last_error = resp
                if attempt < constants.LLM_MAX_RETRIES:
                    log.warning(f"Claude rate limited the request. Retrying in {wait_seconds}s ({attempt}/{constants.LLM_MAX_RETRIES})...")
                    time.sleep(wait_seconds)
                    continue
            resp.raise_for_status()
            return resp.json()
        except requests.exceptions.RequestException as e:
            last_error = e.response if getattr(e, "response", None) is not None else e
            if attempt < constants.LLM_MAX_RETRIES:
                log.warning(f"Claude request failed. Retrying in {constants.LLM_RETRY_SECONDS}s ({attempt}/{constants.LLM_MAX_RETRIES}): {e}")
                time.sleep(constants.LLM_RETRY_SECONDS)
                continue
            break

    if getattr(last_error, "status_code", None) == 429:
        detail = ""
        try:
            detail = last_error.json().get("error", {}).get("message", "")
        except ValueError:
            detail = last_error.text[:300]
        raise SystemExit(
            "Claude returned 429 Too Many Requests after retries. "
            "This usually means rate limit or insufficient quota. "
            "Wait a minute and retry, or check billing/limits for the API key. "
            f"Details: {detail}"
        )

    raise SystemExit(f"Claude request failed after retries: {last_error}")


def claude_text(data: dict) -> str:
    chunks = []
    for item in data.get("content", []):
        if item.get("type") == "text" and item.get("text"):
            chunks.append(item["text"])
    return "\n".join(chunks).strip()


def post_gemini_with_retries(payload: dict) -> dict:
    last_error = None
    url = constants.GEMINI_GENERATE_CONTENT_URL.format(model=constants.GEMINI_MODEL)
    params = {"key": constants.GEMINI_API_KEY}
    for attempt in range(1, constants.LLM_MAX_RETRIES + 1):
        try:
            resp = requests.post(
                url,
                params=params,
                headers={"Content-Type": "application/json"},
                json=payload,
                timeout=constants.LLM_TIMEOUT,
            )
            if resp.status_code == 429:
                retry_after = resp.headers.get("retry-after")
                wait_seconds = int(retry_after) if retry_after and retry_after.isdigit() else constants.LLM_RETRY_SECONDS
                last_error = resp
                if attempt < constants.LLM_MAX_RETRIES:
                    log.warning(f"Gemini rate limited the request. Retrying in {wait_seconds}s ({attempt}/{constants.LLM_MAX_RETRIES})...")
                    time.sleep(wait_seconds)
                    continue
            resp.raise_for_status()
            return resp.json()
        except requests.exceptions.RequestException as e:
            last_error = e.response if getattr(e, "response", None) is not None else e
            if attempt < constants.LLM_MAX_RETRIES:
                log.warning(f"Gemini request failed. Retrying in {constants.LLM_RETRY_SECONDS}s ({attempt}/{constants.LLM_MAX_RETRIES}): {e}")
                time.sleep(constants.LLM_RETRY_SECONDS)
                continue
            break

    if getattr(last_error, "status_code", None) == 429:
        detail = ""
        try:
            detail = last_error.json().get("error", {}).get("message", "")
        except ValueError:
            detail = last_error.text[:300]
        raise SystemExit(
            "Gemini returned 429 Too Many Requests after retries. "
            "This usually means rate limit or insufficient quota. "
            "Wait a minute and retry, or check billing/limits for the API key. "
            f"Details: {detail}"
        )

    raise SystemExit(f"Gemini request failed after retries: {last_error}")


def gemini_text(data: dict) -> str:
    chunks = []
    for candidate in data.get("candidates", []):
        content = candidate.get("content", {})
        for part in content.get("parts", []):
            text = part.get("text")
            if text:
                chunks.append(text)
    return "\n".join(chunks).strip()


def call_llm_text(system_prompt: str, user_prompt: str, max_tokens: int | None = None) -> str:
    max_tokens = max_tokens or constants.LLM_MAX_OUTPUT_TOKENS
    if constants.LLM_PROVIDER == "claude":
        payload = {
            "model": constants.CLAUDE_MODEL,
            "max_tokens": max_tokens,
            "temperature": 0.2,
            "system": system_prompt,
            "messages": [{"role": "user", "content": user_prompt}],
        }
        return claude_text(post_claude_with_retries(payload))

    if constants.LLM_PROVIDER == "gemini":
        payload = {
            "contents": [
                {
                    "role": "user",
                    "parts": [{"text": system_prompt + "\n\n" + user_prompt}],
                }
            ],
            "generationConfig": {
                "temperature": 0.2,
                "maxOutputTokens": max_tokens,
                "responseMimeType": "application/json",
            },
        }
        return gemini_text(post_gemini_with_retries(payload))

    payload = {
        "model": constants.OPENAI_MODEL,
        "temperature": 0.2,
        "response_format": {"type": "json_object"},
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    }
    data = post_openai_with_retries(payload)
    return data["choices"][0]["message"]["content"]


def repair_json_response(bad_content: str) -> dict:
    log.warning("LLM returned malformed JSON. Asking provider to repair the JSON response...")
    repaired = call_llm_text(
        "You repair malformed JSON. Return only valid JSON and do not add new facts.",
        (
            "Fix the following malformed resume JSON so it exactly follows this schema. "
            "Do not invent missing facts; use empty strings or empty lists when needed.\n\n"
            f"Schema:\n{resume_json_schema_text()}\n\n"
            f"Malformed JSON:\n{bad_content[:20000]}"
        ),
        max_tokens=constants.LLM_MAX_OUTPUT_TOKENS,
    )
    return parse_json_object(repaired)


def tailor_resume(base_resume: str, job_description: str) -> dict:
    if constants.LLM_PROVIDER not in {"openai", "claude", "gemini"}:
        raise SystemExit("LLM_PROVIDER must be 'openai', 'claude', or 'gemini'")
    if constants.LLM_PROVIDER == "openai" and not constants.OPENAI_API_KEY:
        raise SystemExit("OPENAI_API_KEY is not set in .env")
    if constants.LLM_PROVIDER == "claude" and not constants.CLAUDE_API_KEY:
        raise SystemExit("CLAUDE_API_KEY or ANTHROPIC_API_KEY is not set in .env")
    if constants.LLM_PROVIDER == "gemini" and not constants.GEMINI_API_KEY:
        raise SystemExit("GEMINI_API_KEY is not set in .env")

    prompts = load_prompts()
    user_content = render_prompt(
        prompts["tailor_resume_user"],
        base_resume=base_resume[:constants.MAX_RESUME_CHARS],
        job_description=job_description[:constants.MAX_JOB_TEXT_CHARS],
    )

    content = call_llm_text(
        prompts["tailor_resume_system"],
        user_content,
        max_tokens=constants.LLM_MAX_OUTPUT_TOKENS,
    )

    if not content:
        raise SystemExit(f"{constants.LLM_PROVIDER} returned an empty response")
    try:
        return parse_json_object(content)
    except json.JSONDecodeError:
        return repair_json_response(content)


def add_bullets(document: Document, bullets: list[str]):
    for bullet in bullets:
        cleaned = clean_bullet_text(bullet)
        if cleaned:
            document.add_paragraph(cleaned, style="List Bullet")


def add_section_heading(document: Document, heading: str):
    paragraph = document.add_paragraph()
    if Pt:
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(heading)
    run.bold = True
    if Pt:
        run.font.size = Pt(12)


def add_bold_paragraph(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.add_run(text).bold = True

def add_labeled_paragraph(document: Document, label: str, text: str):
    paragraph = document.add_paragraph()
    paragraph.add_run(label).bold = True
    if text:
        paragraph.add_run(text)


def as_list(value: Any) -> list:
    return value if isinstance(value, list) else []


def clean_bullet_text(value: object) -> str:
    text = str(value).strip()
    return re.sub(r"^\s*(?:[-*•‣–—]|\d+[.)])\s+", "", text).strip()


def write_resume_docx(result: dict, output_path: Path):
    if Document is None:
        write_resume_docx_basic(result, output_path)
        return

    resume = result.get("resume")
    if not isinstance(resume, dict):
        raise SystemExit("LLM response did not include resume object")

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Calibri"

    name = str(resume.get("name") or "Tailored Resume").strip()
    document.add_heading(name, level=0)

    headline = str(resume.get("headline") or "").strip()
    if headline:
        document.add_paragraph(headline)

    contact = " | ".join(str(item).strip() for item in as_list(resume.get("contact")) if str(item).strip())
    if contact:
        document.add_paragraph(contact)

    summary = as_list(resume.get("summary"))
    if summary:
        add_section_heading(document, "Summary")
        add_bullets(document, summary)

    skills = resume.get("skills")
    if isinstance(skills, dict) and skills:
        add_section_heading(document, "Skills")
        for category, values in skills.items():
            items = ", ".join(str(item).strip() for item in as_list(values) if str(item).strip())
            if items:
                add_labeled_paragraph(document, f"{category}: ", items)

    experience = as_list(resume.get("experience"))
    if experience:
        add_section_heading(document, "Experience")
        for role in experience:
            if not isinstance(role, dict):
                continue
            title_line = " - ".join(
                part for part in [
                    str(role.get("title") or "").strip(),
                    str(role.get("company") or "").strip(),
                ] if part
            )
            if title_line:
                add_bold_paragraph(document, title_line)
            meta = " | ".join(
                part for part in [
                    str(role.get("location") or "").strip(),
                    str(role.get("dates") or "").strip(),
                ] if part
            )
            if meta:
                document.add_paragraph(meta)
            add_bullets(document, as_list(role.get("bullets")))

    projects = as_list(resume.get("projects"))
    if projects:
        add_section_heading(document, "Projects")
        for project in projects:
            if not isinstance(project, dict):
                continue
            name = str(project.get("name") or "").strip()
            if name:
                add_bold_paragraph(document, name)
            add_bullets(document, as_list(project.get("bullets")))

    education = as_list(resume.get("education"))
    if education:
        add_section_heading(document, "Education")
        for edu in education:
            if not isinstance(edu, dict):
                continue
            line = " - ".join(
                part for part in [
                    str(edu.get("degree") or "").strip(),
                    str(edu.get("school") or "").strip(),
                ] if part
            )
            if line:
                add_bold_paragraph(document, line)
            details = str(edu.get("details") or "").strip()
            if details:
                document.add_paragraph(details)

    certifications = as_list(resume.get("certifications"))
    if certifications:
        add_section_heading(document, "Certifications")
        add_bullets(document, certifications)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def docx_paragraph(text: str, style: str | None = None, bold: bool = False) -> str:
    style_xml = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style else ""
    bold_xml = "<w:rPr><w:b/></w:rPr>" if bold else ""
    return (
        "<w:p>"
        f"{style_xml}"
        "<w:r>"
        f"{bold_xml}"
        f"<w:t>{escape(text)}</w:t>"
        "</w:r>"
        "</w:p>"
    )

def docx_labeled_paragraph(label: str, text: str, style: str | None = None) -> str:
    style_xml = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style else ""
    return (
        "<w:p>"
        f"{style_xml}"
        "<w:r><w:rPr><w:b/></w:rPr>"
        f"<w:t>{escape(label)}</w:t>"
        "</w:r>"
        "<w:r>"
        f"<w:t>{escape(text)}</w:t>"
        "</w:r>"
        "</w:p>"
    )


def flatten_resume_lines(resume: dict) -> list[tuple[str, str | None, bool, str | None]]:
    lines = []
    lines.append((str(resume.get("name") or "Tailored Resume").strip(), "Title", True, None))
    headline = str(resume.get("headline") or "").strip()
    if headline:
        lines.append((headline, None, False, None))
    contact = " | ".join(str(item).strip() for item in as_list(resume.get("contact")) if str(item).strip())
    if contact:
        lines.append((contact, None, False, None))

    def heading(text: str):
        lines.append((text, "Heading1", True, None))

    summary = as_list(resume.get("summary"))
    if summary:
        heading("Summary")
        for item in summary:
            cleaned = clean_bullet_text(item)
            if cleaned:
                lines.append((f"• {cleaned}", None, False, None))

    skills = resume.get("skills")
    if isinstance(skills, dict) and skills:
        heading("Skills")
        for category, values in skills.items():
            items = ", ".join(str(item).strip() for item in as_list(values) if str(item).strip())
            if items:
                lines.append((f"{category}: ", None, False, items))

    experience = as_list(resume.get("experience"))
    if experience:
        heading("Experience")
        for role in experience:
            if not isinstance(role, dict):
                continue
            title_line = " - ".join(
                part for part in [
                    str(role.get("title") or "").strip(),
                    str(role.get("company") or "").strip(),
                ] if part
            )
            if title_line:
                lines.append((title_line, None, True, None))
            meta = " | ".join(
                part for part in [
                    str(role.get("location") or "").strip(),
                    str(role.get("dates") or "").strip(),
                ] if part
            )
            if meta:
                lines.append((meta, None, False, None))
            for bullet in as_list(role.get("bullets")):
                cleaned = clean_bullet_text(bullet)
                if cleaned:
                    lines.append((f"• {cleaned}", None, False, None))

    projects = as_list(resume.get("projects"))
    if projects:
        heading("Projects")
        for project in projects:
            if not isinstance(project, dict):
                continue
            name = str(project.get("name") or "").strip()
            if name:
                lines.append((name, None, True, None))
            for bullet in as_list(project.get("bullets")):
                cleaned = clean_bullet_text(bullet)
                if cleaned:
                    lines.append((f"• {cleaned}", None, False, None))

    education = as_list(resume.get("education"))
    if education:
        heading("Education")
        for edu in education:
            if not isinstance(edu, dict):
                continue
            line = " - ".join(
                part for part in [
                    str(edu.get("degree") or "").strip(),
                    str(edu.get("school") or "").strip(),
                ] if part
            )
            if line:
                lines.append((line, None, True, None))
            details = str(edu.get("details") or "").strip()
            if details:
                lines.append((details, None, False, None))

    certifications = as_list(resume.get("certifications"))
    if certifications:
        heading("Certifications")
        for item in certifications:
            cleaned = clean_bullet_text(item)
            if cleaned:
                lines.append((f"• {cleaned}", None, False, None))

    return [(text, style, bold, tail) for text, style, bold, tail in lines if text]


def write_resume_docx_basic(result: dict, output_path: Path):
    resume = result.get("resume")
    if not isinstance(resume, dict):
        raise SystemExit("LLM response did not include resume object")

    body = "\n".join(
        docx_labeled_paragraph(text, tail, style) if tail is not None else docx_paragraph(text, style, bold)
        for text, style, bold, tail in flatten_resume_lines(resume)
    )
    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    {body}
    <w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="720" w:right="720" w:bottom="720" w:left="720"/></w:sectPr>
  </w:body>
</w:document>"""
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", rels)
        archive.writestr("word/document.xml", document_xml)


def write_outputs(result: dict, output_path: Path, analysis_path: Path):
    output_path.parent.mkdir(parents=True, exist_ok=True)
    analysis_path.parent.mkdir(parents=True, exist_ok=True)

    write_resume_docx(result, output_path)
    analysis_path.write_text(json.dumps(result, indent=2), encoding="utf-8")


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Tailor a base resume to a job URL and write a DOCX resume.")
    parser.add_argument(
        "--resume",
        default=str(constants.BASE_RESUME_PATH),
        help="Path to base resume: .txt, .md, or .docx. Defaults to BASE_RESUME_PATH.",
    )
    parser.add_argument("--job-url", help="Job posting URL to extract. If omitted, the script prompts for it.")
    parser.add_argument(
        "--output",
        help="DOCX resume output path. Defaults to ABHIRAM_role_name.docx in outputs/.",
    )
    parser.add_argument(
        "--analysis-output",
        help="Full JSON analysis output path. Defaults beside the DOCX output.",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Load inputs and print character counts without calling the LLM",
    )
    return parser


def main():
    args = build_parser().parse_args()
    resume_path = resolve_resume_path(args.resume)
    job_url = args.job_url.strip() if args.job_url else prompt_for_job_url()
    base_resume = read_text_file(resume_path)[:constants.MAX_RESUME_CHARS]
    job_description = load_job_description(job_url)

    if not base_resume:
        raise SystemExit("Base resume is empty.")
    if not job_description:
        raise SystemExit("Job description is empty.")

    if args.dry_run:
        log.info(f"Loaded resume: {len(base_resume)} characters")
        log.info(f"Loaded job description: {len(job_description)} characters")
        return

    result = tailor_resume(base_resume, job_description)
    default_docx, default_analysis = default_output_paths(result, job_description)
    output_path = Path(args.output).expanduser() if args.output else default_docx
    analysis_path = Path(args.analysis_output).expanduser() if args.analysis_output else default_analysis
    write_outputs(result, output_path, analysis_path)

    log.info(f"Tailored resume saved to {output_path}")
    log.info(f"Analysis saved to {analysis_path}")


if __name__ == "__main__":
    main()
